"""
report_builder.py — Render the final analytics report as a Word document.

WHY THIS EXISTS:
  This is the project's "generates an output" deliverable. It assembles:
    • An executive-summary narrative (from insight_generator)
    • KPI tables (from kpi_engine)
    • Charts (matplotlib PNGs embedded)
    • A knowledge-graph snapshot (top connected entities)
    • Full source citations (filename + page) for auditability

  The output is a polished .docx an analyst could send to a stakeholder —
  turning a pile of PDFs into a single decision-ready report.

WHY python-docx + matplotlib:
  python-docx produces native Word files (editable, shareable in business
  settings). matplotlib renders charts to PNG which docx embeds. Both are
  pure-Python and need no external services.
"""

from __future__ import annotations
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")  # headless
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.analytics.kpi_engine import KPISummary
from app.config import settings
from app.utils.logger import logger

BLUE = RGBColor(0x1F, 0x4E, 0x79)


def _bar_chart(data: dict, title: str, path: Path, xlabel: str = "") -> Path | None:
    if not data:
        return None
    fig, ax = plt.subplots(figsize=(6.5, 3.2))
    keys = [str(k)[:22] for k in data.keys()]
    ax.barh(keys, list(data.values()), color="#2E5C8A")
    ax.set_title(title, fontsize=12, fontweight="bold", color="#1F4E79")
    ax.set_xlabel(xlabel)
    ax.invert_yaxis()
    fig.tight_layout()
    fig.savefig(path, dpi=130, bbox_inches="tight")
    plt.close(fig)
    return path


def build_report(
    narrative: str,
    kpi: KPISummary,
    graph_stats: dict,
    source_files: list[str],
) -> Path:
    """Assemble and save the analytics report. Returns the .docx path."""
    doc = Document()

    # ── Styling defaults ──
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Financial Document Intelligence Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run(f"Auto-generated from {len(source_files)} document(s) · "
                    f"{datetime.now():%d %b %Y}")
    s.italic = True
    s.font.size = Pt(10)
    s.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()

    # ── Executive summary ──
    h = doc.add_heading("Executive Summary", level=1)
    h.runs[0].font.color.rgb = BLUE
    doc.add_paragraph(narrative)

    # ── Key metrics table ──
    if kpi.total_by_metric:
        h = doc.add_heading("Key Financial Metrics", level=1)
        h.runs[0].font.color.rgb = BLUE
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Total Value"
        for metric, total in kpi.total_by_metric.items():
            row = table.add_row().cells
            row[0].text = metric.title()
            row[1].text = f"{total:,.0f}"

        # Chart
        chart_path = settings.reports_dir / "_metric_chart.png"
        if _bar_chart(kpi.total_by_metric, "Total by Metric Type", chart_path, "Value"):
            doc.add_paragraph()
            doc.add_picture(str(chart_path), width=Inches(6))

    # ── Top entities ──
    if kpi.top_entities:
        h = doc.add_heading("Top Entities by Financial Significance", level=1)
        h.runs[0].font.color.rgb = BLUE
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Entity"
        table.rows[0].cells[1].text = "Aggregate Value"
        for e in kpi.top_entities:
            row = table.add_row().cells
            row[0].text = str(e["entity"])
            row[1].text = f"{e['total_value']:,.0f}"

        chart_data = {e["entity"]: e["total_value"] for e in kpi.top_entities[:8]}
        chart_path = settings.reports_dir / "_entity_chart.png"
        if _bar_chart(chart_data, "Top Entities by Value", chart_path, "Value"):
            doc.add_paragraph()
            doc.add_picture(str(chart_path), width=Inches(6))

    # ── Knowledge graph snapshot ──
    h = doc.add_heading("Knowledge Graph Overview", level=1)
    h.runs[0].font.color.rgb = BLUE
    doc.add_paragraph(
        f"The system constructed a knowledge graph of "
        f"{graph_stats.get('total_entities', 0)} entities connected by "
        f"{graph_stats.get('total_relationships', 0)} relationships extracted "
        f"across the document set."
    )
    top_connected = graph_stats.get("top_connected", [])
    if top_connected:
        doc.add_paragraph("Most connected entities (network hubs):")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Entity", "Type", "Connections"
        for tc in top_connected:
            row = table.add_row().cells
            row[0].text = str(tc.get("name", ""))
            row[1].text = str(tc.get("type", ""))
            row[2].text = str(tc.get("degree", ""))

    # ── Derived insights ──
    if kpi.insights:
        h = doc.add_heading("Derived Insights", level=1)
        h.runs[0].font.color.rgb = BLUE
        for ins in kpi.insights:
            doc.add_paragraph(ins, style="List Bullet")

    # ── Sources ──
    h = doc.add_heading("Source Documents", level=1)
    h.runs[0].font.color.rgb = BLUE
    for f in source_files:
        doc.add_paragraph(f, style="List Bullet")

    # ── Save ──
    out_path = settings.reports_dir / f"report_{datetime.now():%Y%m%d_%H%M%S}.docx"
    doc.save(str(out_path))
    logger.success(f"Report saved → {out_path}")
    return out_path
